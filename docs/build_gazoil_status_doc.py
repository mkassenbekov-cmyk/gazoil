from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "GAZOIL_Текущее_состояние_и_дорожная_карта.docx"
LOGO = ROOT / "gazoil-logo.jpg"

BLUE = "006FD6"
DARK_BLUE = "18324A"
LIGHT_BLUE = "EAF4FF"
PALE_BLUE = "F4F8FC"
GREEN = "248A3D"
LIGHT_GREEN = "E8F7ED"
AMBER = "BF5B00"
LIGHT_AMBER = "FFF2DF"
RED = "C62828"
LIGHT_RED = "FDEBEC"
INK = "1D1D1F"
MUTED = "667085"
LINE = "D8DEE8"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=INK, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    existing_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId"))
    ]
    next_abstract = max(existing_ids or [0]) + 1

    def add_abstract(abstract_id, num_fmt, level_text):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        lvl.append(text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(abstract_id):
        nums = [
            int(el.get(qn("w:numId")))
            for el in numbering.findall(qn("w:num"))
            if el.get(qn("w:numId"))
        ]
        num_id = max(nums or [0]) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)
        return num_id

    add_abstract(next_abstract, "bullet", "•")
    bullet_num_id = add_num(next_abstract)
    add_abstract(next_abstract + 1, "decimal", "%1.")
    decimal_num_id = add_num(next_abstract + 1)
    return bullet_num_id, decimal_num_id


def duplicate_numbering_instance(doc, source_num_id):
    numbering = doc.part.numbering_part.element
    source = next(
        el for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) == str(source_num_id)
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId"))
    ]
    num_id = max(existing_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), abstract_id)
    num.append(abstract_num_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p._p.get_or_add_pPr().append(num_pr)
    set_paragraph_spacing(p, after=4)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    set_paragraph_spacing(
        p,
        before={1: 18, 2: 14, 3: 10}[level],
        after={1: 10, 2: 7, 3: 5}[level],
        line=1.15,
    )
    return p


def add_section_heading(doc, text):
    p = add_heading(doc, text, 1)
    p.paragraph_format.page_break_before = True
    return p


def add_body(doc, text, bold_lead=None, color=INK, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=10.5, bold=True, color=color)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.5, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=color)
    return p


def add_callout(doc, label, text, tone="blue"):
    colors = {
        "blue": (LIGHT_BLUE, BLUE),
        "green": (LIGHT_GREEN, GREEN),
        "amber": (LIGHT_AMBER, AMBER),
        "red": (LIGHT_RED, RED),
        "gray": (PALE_BLUE, DARK_BLUE),
    }
    fill, accent = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.18)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, after=0, line=1.05)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=DARK_BLUE)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.08)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_status_row(doc, title, status, detail, status_color):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2350, 1500, 5510])
    values = [title, status, detail]
    for idx, value in enumerate(values):
        cell = table.cell(0, idx)
        set_cell_shading(cell, PALE_BLUE if idx != 1 else status_color[0])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.08)
        run = p.add_run(value)
        set_run_font(run, size=9.5, bold=(idx < 2), color=status_color[1] if idx == 1 else INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else DARK_BLUE)
        style.font.size = Pt({1: 16, 2: 13, 3: 11.5}[level])
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
        style.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}[level])


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=0)
    run = hp.add_run("GAZOIL  |  Единая система управления")
    set_run_font(run, size=9, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def cover_page(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(50)
        p.paragraph_format.space_after = Pt(24)
        p.add_run().add_picture(str(LOGO), width=Inches(1.25))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=14)
    run = kicker.add_run("АРХИТЕКТУРА И ТЕКУЩЕЕ СОСТОЯНИЕ ПЛАТФОРМЫ")
    set_run_font(run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=10, line=1.08)
    run = title.add_run("Единая система управления GAZOIL")
    set_run_font(run, size=28, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=42, line=1.2)
    run = subtitle.add_run("Что уже реализовано, как связаны процессы и что делать дальше")
    set_run_font(run, size=15, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2400, 6960])
    metadata = [
        ("Версия документа", "1.0"),
        ("Актуально на", "21 июня 2026 года"),
        ("Состояние проекта", "Рабочий локальный прототип; серверные интеграции ещё не подключены"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_shading(row.cells[0], "E8EEF5")
        set_cell_shading(row.cells[1], "F8FAFC")
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0)
            run = p.add_run(text)
            set_run_font(run, size=10, bold=(idx == 0), color=DARK_BLUE if idx == 0 else INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Внутренний рабочий документ GAZOIL")
    set_run_font(run, size=9.5, italic=True, color=MUTED)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_num_id, decimal_num_id = add_numbering_definitions(doc)
    cover_page(doc)

    add_heading(doc, "Как пользоваться этим документом", 1)
    add_callout(
        doc,
        "Главная мысль",
        "Бизнес-логика и большинство экранов уже работают в локальном прототипе. "
        "Но это ещё не промышленная система: нет серверной базы данных, реальных интеграций и централизованного файлового хранилища.",
        "blue",
    )
    add_body(
        doc,
        "Документ нужен как единая карта проекта: он фиксирует уже реализованные роли и маршруты, "
        "показывает границы прототипа и помогает принимать дальнейшие решения в правильном порядке.",
    )

    add_heading(doc, "Обозначения статуса", 2)
    add_status_row(doc, "Работает", "ГОТОВО", "Можно проверить в интерфейсе локального приложения.", (LIGHT_GREEN, GREEN))
    add_status_row(doc, "Демо-реализация", "MOCK", "Логика готова, но внешний сервис пока имитируется.", (LIGHT_AMBER, AMBER))
    add_status_row(doc, "Следующий этап", "НУЖНО", "Требует сервера, интеграции, данных или управленческого решения.", (LIGHT_RED, RED))

    add_heading(doc, "Содержание", 2)
    contents = [
        "1. Краткое резюме текущего состояния",
        "2. Общая сквозная логика платформы",
        "3. Компании и маршрутизация",
        "4. Пользователи, роли и кабинеты",
        "5. Модули платформы",
        "6. Подробная логика процессов",
        "7. Задачи, документы, уведомления и журнал",
        "8. ИИ-помощник и корпоративные сообщения",
        "9. Техническое состояние и ограничения",
        "10. Что необходимо решить дальше",
        "11. Рекомендуемая дорожная карта",
        "12. Контрольный список следующего этапа",
        "Приложение: стартовые учётные записи",
    ]
    for item in contents:
        p = add_body(doc, item, after=1)
        p.paragraph_format.left_indent = Inches(0.15)

    add_section_heading(doc, "1. Краткое резюме текущего состояния")
    add_body(
        doc,
        "На текущий момент создан связанный локальный прототип внутренней системы управления коммерческими процессами GAZOIL. "
        "В нём уже можно входить под разными сотрудниками, видеть разные кабинеты, создавать и вести карточки процессов, "
        "проходить согласования, формировать mock-счета 1С, ставить задачи, переписываться и использовать персонального ИИ-помощника.",
    )
    add_heading(doc, "Что уже работает", 2)
    ready_items = [
        "Авторизация, стартовые пользователи, роли и ограничения доступа.",
        "Раздельный доступ по ТОО «GAZOIL TRADE» и ТОО «УГХ GAZOIL».",
        "Клиенты, обращения, договоры, заказы, продления, возвраты и тендеры.",
        "Статусы процессов, проверки переходов и связанные карточки.",
        "Согласования директоров с решением и обязательным комментарием при возврате.",
        "Задачи, документы, уведомления, отчёты и журнал действий.",
        "Отдельные кабинеты коммерческого директора, директоров компаний, учредителя и сотрудников.",
        "CRM входящих обращений для менеджеров обслуживания: WhatsApp, телефония и email.",
        "Персональный ИИ-помощник и корпоративные личные/групповые сообщения.",
        "Казахский и русский интерфейс, логотип и единый визуальный стиль.",
    ]
    for item in ready_items:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "Что пока не является промышленной интеграцией", 2)
    mock_items = [
        "1С работает в mock-режиме: счёт и подтверждение оплаты создаются внутри прототипа.",
        "WhatsApp, телефония и email представлены как единая очередь, но реальные каналы ещё не подключены.",
        "ИИ-помощник использует локальную ролевую логику, а не внешнюю языковую модель.",
        "Сообщения сохраняются в браузере, а не на сервере.",
        "Пароли хэшируются на стороне браузера, что подходит только для прототипа.",
        "Файлы документов пока представлены метаданными; реального хранилища файлов нет.",
    ]
    for item in mock_items:
        add_list_item(doc, item, bullet_num_id)

    add_callout(
        doc,
        "Статус кода",
        "Основной объём платформы отправлен в GitHub в ветку codex/bitrix-process-architecture, коммит fff1233. "
        "Последняя корректировка CRM-кабинета менеджеров обслуживания пока находится локально и должна быть отправлена отдельным коммитом после подтверждения.",
        "amber",
    )

    add_section_heading(doc, "2. Общая сквозная логика платформы")
    add_callout(
        doc,
        "Основная цепочка",
        "Клиент → входящее обращение → классификация → договор / заказ / продление / возврат / тендер → "
        "финансовые и директорские проверки → исполнение → документы → закрытие → отчёты и контроль.",
        "green",
    )
    flow_steps = [
        "Клиент обращается через WhatsApp, телефон, email, офис, сайт или другой канал.",
        "Обращение попадает конкретному менеджеру обслуживания: Ельжану, Жанаре или Екатерине.",
        "Менеджер отвечает клиенту, определяет тему и фиксирует данные в карточке.",
        "Из обращения создаётся связанный профильный процесс без повторного ввода основных данных.",
        "Процесс идёт по маршруту выбранной компании и назначается ответственным подразделениям.",
        "Если требуется решение директора, карточка автоматически появляется у директора своей компании.",
        "Финансовые факты должны приходить из 1С; в прототипе это имитируется mock-кнопками.",
        "Все важные действия фиксируются в истории карточки, уведомлениях и audit log.",
        "Руководители и учредитель видят только свою аналитику и сигналы в соответствии с ролью.",
    ]
    flow_num_id = duplicate_numbering_instance(doc, decimal_num_id)
    for item in flow_steps:
        add_list_item(doc, item, flow_num_id)

    add_heading(doc, "Принцип единой карточки", 2)
    add_body(
        doc,
        "Каждая операция должна существовать в системе как отдельная карточка с собственным ID, компанией, клиентом, "
        "ответственным, стадией, сроком, документами, задачами, комментариями и историей. "
        "Ручное переключение стадий заменено контролируемыми переходами с обязательными условиями.",
    )

    add_heading(doc, "Принцип ролевой видимости", 2)
    add_body(
        doc,
        "Пользователь видит не всю систему, а только разрешённый ему рабочий контур. "
        "Менеджер обслуживания видит назначенные обращения и процессы; бухгалтер — финансовые карточки своей компании; "
        "директор — процессы своей компании; учредитель — только агрегированную аналитику и просрочки директоров.",
    )

    add_section_heading(doc, "3. Компании и маршрутизация")
    add_table(
        doc,
        ["Компания", "Направление", "Директор", "Бухгалтер", "Основной маршрут"],
        [
            ["ТОО «GAZOIL TRADE»", "ГСМ", "Рашид / Хуснутдинов Р. Р.", "Татьяна", "ГСМ → Татьяна → Рашид"],
            ["ТОО «УГХ GAZOIL»", "Газ", "Арсен / Киікбай А. Б.", "Айнура", "Газ → Айнура → Арсен"],
        ],
        [2100, 900, 2200, 1300, 2860],
        font_size=8.8,
    )
    add_heading(doc, "Правило поля «Наша компания»", 2)
    rules = [
        "Поле обязательно для каждого процесса.",
        "От него зависят директор, бухгалтер, отчёты, права доступа и маршрут согласования.",
        "После выхода со стартовой стадии компания блокируется от обычного изменения.",
        "Смена компании после начала согласования должна быть доступна только администратору с обязательной причиной — это правило частично заложено и требует серверного усиления.",
    ]
    for item in rules:
        add_list_item(doc, item, bullet_num_id)

    add_callout(
        doc,
        "Пример",
        "Если возврат создан по УГХ GAZOIL, финансовая проверка идёт Айнуре, а директорское решение — Арсену. "
        "Если возврат создан по GAZOIL TRADE, проверка идёт Татьяне, решение — Рашиду.",
        "blue",
    )

    add_section_heading(doc, "4. Пользователи, роли и кабинеты")
    user_rows = [
        ["Мади", "Коммерческий директор / администратор", "Обе компании; все процессы; пользователи; отчёты; audit log"],
        ["Рашид", "Директор GAZOIL TRADE", "Только GAZOIL TRADE; согласования, продажи и отчёты компании"],
        ["Арсен", "Директор УГХ GAZOIL", "Только УГХ GAZOIL; согласования, продажи и отчёты компании"],
        ["Диана", "Старший / главный менеджер", "Контроль команды, сроков, просрочек и качества выполнения"],
        ["Ельжан", "Менеджер обслуживания", "Личная CRM-очередь входящих и назначенные процессы"],
        ["Жанара", "Менеджер обслуживания", "Личная CRM-очередь входящих и назначенные процессы"],
        ["Екатерина", "Менеджер обслуживания", "Личная CRM-очередь входящих и назначенные процессы"],
        ["Ольга", "Госзакупщик", "Тендеры, госорганизации, продления и возвраты госорганизаций"],
        ["Павел", "Менеджер продаж", "Новые клиенты, обращения, первичные продажи и договоры"],
        ["Ильяс", "Менеджер опта", "Оптовые заявки, договоры, счета и контроль оплаты"],
        ["Татьяна", "Бухгалтер GAZOIL TRADE", "Финансовые процессы только GAZOIL TRADE"],
        ["Айнура", "Бухгалтер УГХ GAZOIL", "Финансовые процессы только УГХ GAZOIL"],
        ["Учредитель", "Наблюдатель", "Продажи, динамика, риски, отчёты и просрочки директоров; без редактирования"],
    ]
    add_table(doc, ["Пользователь", "Роль", "Рабочий контур"], user_rows, [1500, 2500, 5360], font_size=8.4)

    add_heading(doc, "Кабинет менеджера обслуживания", 2)
    service_features = [
        "Персональная входящая очередь из WhatsApp, телефонии и email.",
        "Фильтр по каналам, канбан стадий и журнал коммуникаций.",
        "Просмотр сообщения клиента и ответ прямо из карточки.",
        "Фиксация первого ответа и перевод обращения в работу.",
        "Создание связанного договора, заказа, продления, возврата или тендера.",
        "Менеджер не распределяет обращения между коллегами и не согласовывает решения директора.",
    ]
    for item in service_features:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "Кабинет Дианы", 2)
    add_body(
        doc,
        "Диана не должна обрабатывать входящие заявки вместо менеджеров. Её задача — видеть загрузку Ельжана, Жанары и Екатерины, "
        "контролировать просрочки, качество карточек и соблюдение нормативов. Значения SLA ещё необходимо утвердить отдельно.",
    )

    add_heading(doc, "Дополнительные роли, предусмотренные моделью", 2)
    add_body(
        doc,
        "В RBAC уже предусмотрены роли кассира, бухгалтера по банку и системного администратора. "
        "Однако реальные сотрудники, логины и окончательные обязанности по этим ролям ещё не утверждены.",
    )

    add_section_heading(doc, "5. Модули платформы")
    module_rows = [
        ["Главная панель", "Работает", "Ролевые показатели, очереди, сигналы и аналитика"],
        ["Клиенты", "Работает", "Единая карточка, предупреждение о дубле БИН, связанные процессы"],
        ["Обращения", "Работает", "Омниканальная демо-очередь, классификация и преобразование"],
        ["Договоры", "Работает", "Статусы, подписанный файл, контроль окончания"],
        ["Заказы и выдача", "Работает + mock", "Договор, счёт, оплата 1С, доверенность, выдача, ЭСФ"],
        ["Продления", "Работает", "Решение директора, перерасчёт, касса, переход в возврат"],
        ["Возвраты", "Работает", "Аннулирование, бухгалтерия, директор, банк, платёжное поручение"],
        ["Тендеры", "Работает + mock", "Импорт JSON/API, согласование, подача, результат, договор"],
        ["Задачи", "Работает", "Исполнитель, срок, статус, переназначение, поручения учредителя"],
        ["Документы", "Работает как реестр", "Метаданные есть; реальное файловое хранилище ещё не подключено"],
        ["Отчёты", "Работает", "Ролевые сводки и выбор периода"],
        ["Пользователи и роли", "Работает", "Создание пользователя, роль, отключение, сброс пароля"],
        ["Журнал действий", "Работает", "Входы, статусы, решения, документы, пользователи и интеграции"],
        ["ИИ-помощник", "Локальная демо-логика", "Персональная история и ролевые ответы"],
        ["Сообщения", "Локальная демо-логика", "Личные диалоги, группы и непрочитанные сообщения"],
    ]
    add_table(doc, ["Модуль", "Статус", "Что реализовано"], module_rows, [1850, 1550, 5960], font_size=8.4)

    add_section_heading(doc, "6. Подробная логика процессов")

    add_heading(doc, "6.1. Входящее обращение", 2)
    add_callout(
        doc,
        "Маршрут",
        "Канал → менеджер обслуживания → первый ответ → классификация → связанный профильный процесс → контроль Дианы.",
        "blue",
    )
    appeal_rules = [
        "Источники: WhatsApp, телефония, email, офис, сайт, руководитель и другие.",
        "Закрытие без решения требует обязательной причины.",
        "Из обращения можно создать договор, заказ, продление, возврат или тендер.",
        "В карточке сохраняется ссылка на созданный процесс.",
        "Реальное автоматическое распределение входящих пока не подключено; сейчас используются демо-данные и назначенные владельцы.",
    ]
    for item in appeal_rules:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "6.2. Договор", 2)
    add_callout(
        doc,
        "Маршрут",
        "Реквизиты → подготовка проекта → отправка клиенту → подписание → действует → скоро истекает → перезаключение.",
        "green",
    )
    contract_rules = [
        "Договор нельзя активировать без подписанного файла.",
        "Заказ нельзя создавать без действующего договора по компании и продукту.",
        "За 15 дней до окончания система переводит договор в «Скоро истекает», создаёт задачу и уведомляет Диану и Мади.",
        "Для полноценной автоматизации необходимы реальные даты договоров и серверный планировщик.",
    ]
    for item in contract_rules:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "6.3. Заказ и выдача топлива", 2)
    add_callout(
        doc,
        "Маршрут",
        "Запрос → проверка договора → счёт 1С → ожидание оплаты → подтверждение оплаты → доверенность → выдача → документы → ЭСФ → закрытие.",
        "green",
    )
    order_rules = [
        "Счёт можно сформировать через mock-1С только при наличии действующего договора.",
        "Подтверждение оплаты доступно администратору или бухгалтерской роли.",
        "Без подтверждённой оплаты переход к выдаче запрещён.",
        "Выданный объём не может превышать оплаченный.",
        "Успешное закрытие требует документа выдачи и статуса ЭСФ.",
    ]
    for item in order_rules:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "6.4. Продление талонов и топливных карт", 2)
    add_callout(
        doc,
        "Маршрут",
        "Заявка → документы → директор компании → ответ клиента → коммерческий отдел → касса при необходимости → закрытие.",
        "amber",
    )
    extension_rules = [
        "Решения директора: продлить без перерасчёта, продлить с перерасчётом, отказать или вернуть на доработку.",
        "Для решения с перерасчётом требуется указать сумму.",
        "Без решения директора продление невозможно.",
        "При отказе клиента и выборе возврата создаётся связанная заявка на возврат.",
        "Повторное создание возврата из одной заявки блокируется связью карточек.",
    ]
    for item in extension_rules:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "6.5. Возврат денежных средств", 2)
    add_callout(
        doc,
        "Маршрут",
        "Заявка → документы → аннулирование → бухгалтер компании → акт сверки / письмо → директор компании → бухгалтер по банку → платёжное поручение → закрытие.",
        "amber",
    )
    refund_rules = [
        "Талоны + деньги направляются на аннулирование кассе.",
        "Товарная карта + деньги направляется ответственному за операции с картами.",
        "Только деньги пропускают этап аннулирования.",
        "GAZOIL TRADE проверяет Татьяна; УГХ GAZOIL проверяет Айнура.",
        "Если переплата не подтверждена, возврат закрывается без банковского перечисления.",
        "Если переплата подтверждена, требуется решение директора своей компании.",
        "Успешное закрытие возможно только после номера, даты и файла платёжного поручения.",
    ]
    for item in refund_rules:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "6.6. Тендеры и госзакупки", 2)
    add_callout(
        doc,
        "Маршрут",
        "Импорт / ручное создание → проверка Ольгой → обсуждение → директор компании → участвуем → подготовка → подача → итог → договор.",
        "blue",
    )
    tender_rules = [
        "Предусмотрен mock-импорт тендеров через JSON/API.",
        "Без обязательных данных нельзя отправить карточку директору.",
        "Без решения «Участвуем» нельзя начать подачу заявки.",
        "Подача требует даты, времени и подтверждающего документа.",
        "Успешное закрытие требует номера и файла подписанного договора.",
        "Заложены уведомления за 3 часа, 30 минут и при просрочке дедлайна.",
    ]
    for item in tender_rules:
        add_list_item(doc, item, bullet_num_id)

    add_section_heading(doc, "7. Задачи, документы, уведомления и журнал")
    add_heading(doc, "Задачи", 2)
    task_items = [
        "Задача связана с конкретной карточкой или аналитическим показателем.",
        "Хранятся название, исполнитель, срок, приоритет, статус, создатель и результат.",
        "Диана и Мади могут переназначать задачи.",
        "Учредитель может поставить поручение прямо из показателя продаж или риска.",
        "Назначенный сотрудник видит поручение в своём разделе задач.",
    ]
    for item in task_items:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "Документы", 2)
    add_body(
        doc,
        "В карточках можно регистрировать договоры, счета, акты, накладные, ЭСФ, доверенности, письма клиента, "
        "служебные записки, акты сверки, платёжные поручения и тендерные документы. "
        "Некоторые типы документов автоматически включают соответствующие проверки процесса.",
    )
    add_callout(
        doc,
        "Ограничение",
        "Сейчас фиксируется название и тип документа, но сам файл не загружается в серверное хранилище. "
        "Для промышленной системы нужно подключить объектное хранилище, контроль доступа и версии файлов.",
        "amber",
    )

    add_heading(doc, "Уведомления", 2)
    notification_items = [
        "Новая задача и переназначение.",
        "Просрочка процесса.",
        "Согласование директора.",
        "Оплата подтверждена из 1С.",
        "Договор скоро истекает.",
        "Тендер приближается к дедлайну или просрочен.",
        "Новое личное или групповое сообщение.",
    ]
    for item in notification_items:
        add_list_item(doc, item, bullet_num_id)

    add_heading(doc, "Журнал действий", 2)
    add_body(
        doc,
        "Audit log уже фиксирует входы и выходы, создание и изменение карточек, смену статуса, решения директора, "
        "формирование счёта и подтверждение оплаты, загрузку документов, задачи, пользователей, сообщения и запросы к ИИ.",
    )

    add_section_heading(doc, "8. ИИ-помощник и корпоративные сообщения")
    add_heading(doc, "Персональный ИИ-помощник", 2)
    ai_items = [
        "Доступен каждому авторизованному пользователю.",
        "Имеет отдельную историю для каждого аккаунта.",
        "Отвечает с учётом роли и доступных данных.",
        "Менеджеру может показать задачи и просрочки; директору — согласования; учредителю — продажи и падения.",
        "Не раскрывает финансовую аналитику ролям без соответствующего права.",
    ]
    for item in ai_items:
        add_list_item(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Текущий режим",
        "Ответы формируются локальными правилами. Следующий уровень — подключить реальную AI-модель через сервер, "
        "добавить безопасный поиск по данным и журнал использования.",
        "amber",
    )

    add_heading(doc, "Корпоративные сообщения", 2)
    messages = [
        "Личные диалоги между сотрудниками.",
        "Стартовые группы: «Общий GAZOIL», «Коммерческий отдел», «Руководство», «Бухгалтерия».",
        "Создание новых групп и выбор участников.",
        "Счётчик непрочитанных сообщений и внутренние уведомления.",
        "Новые сотрудники автоматически добавляются в общую группу.",
    ]
    for item in messages:
        add_list_item(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Ограничение",
        "Переписка пока хранится в localStorage конкретного браузера. Для настоящей совместной работы нужен сервер, "
        "WebSocket или другой механизм реального времени, хранение истории и правила удаления/архивации.",
        "red",
    )

    add_section_heading(doc, "9. Техническое состояние и ограничения")
    technical_rows = [
        ["Интерфейс", "HTML, CSS, JavaScript", "Работает локально"],
        ["Хранение данных", "localStorage браузера", "Не подходит для совместной эксплуатации"],
        ["Авторизация", "Хэш SHA-256 в браузере", "Только прототип; нужен серверный Argon2/bcrypt"],
        ["База данных", "Нет", "Нужна PostgreSQL или эквивалент"],
        ["API", "Нет серверного API", "Нужен backend"],
        ["Файлы", "Только метаданные", "Нужно объектное хранилище"],
        ["1С", "Mock", "Нужно согласовать и реализовать реальный обмен"],
        ["WhatsApp / телефония / email", "Демо-очередь", "Нужны реальные провайдеры и webhook"],
        ["ИИ", "Локальные правила", "Нужна серверная AI-интеграция"],
        ["Сообщения", "localStorage", "Нужен сервер и real-time доставка"],
        ["Развёртывание", "Локальный сервер", "Нужны домен, HTTPS, резервные копии и мониторинг"],
    ]
    add_table(doc, ["Слой", "Сейчас", "Что требуется"], technical_rows, [1800, 2500, 5060], font_size=8.6)

    add_heading(doc, "Критические риски при запуске без сервера", 2)
    risks = [
        "Разные пользователи не будут видеть единое актуальное состояние данных.",
        "Очистка браузера может удалить локальные данные.",
        "Нельзя обеспечить надёжную безопасность паролей и прав.",
        "Нет резервного копирования, восстановления и расследования инцидентов.",
        "Нельзя гарантировать доставку сообщений, уведомлений и интеграционных событий.",
        "Невозможно безопасно хранить реальные договоры, платёжные поручения и персональные данные.",
    ]
    for item in risks:
        add_list_item(doc, item, bullet_num_id)

    add_section_heading(doc, "10. Что необходимо решить дальше")
    add_heading(doc, "Управленческие решения", 2)
    decisions = [
        "Утвердить нормативы SLA для входящих: первый ответ, классификация, передача в процесс и закрытие.",
        "Утвердить правила автоматического распределения WhatsApp, звонков и email между Ельжаном, Жанарой и Екатериной.",
        "Назвать конкретного кассира и бухгалтера по банку, определить логины и заместителей.",
        "Подтвердить окончательный список стадий каждого процесса и допустимые возвраты назад.",
        "Определить обязательные поля и документы по каждой стадии.",
        "Определить, кто имеет право отменять карточки, менять компанию и исправлять финансовые данные.",
        "Утвердить набор отчётов для Мади, директоров, Дианы и учредителя.",
    ]
    decision_num_id = duplicate_numbering_instance(doc, decimal_num_id)
    for item in decisions:
        add_list_item(doc, item, decision_num_id)

    add_heading(doc, "Интеграционные решения", 2)
    integrations = [
        "Какая конфигурация 1С используется и какие методы обмена доступны.",
        "Какой WhatsApp-провайдер будет использоваться: официальный Cloud API, Wazzup, Chat2Desk или другой.",
        "Какая телефония используется и может ли она передавать события звонка и записи.",
        "Какая корпоративная почта используется и какие ящики нужно подключить.",
        "Где хранить документы: собственный сервер, S3-совместимое хранилище или облако.",
        "Где будет размещена система: облако, локальный сервер или гибрид.",
    ]
    integration_num_id = duplicate_numbering_instance(doc, decimal_num_id)
    for item in integrations:
        add_list_item(doc, item, integration_num_id)

    add_heading(doc, "Данные, которые нужно подготовить", 2)
    data_items = [
        "Полный список сотрудников, телефонов, email, ролей и замещений.",
        "Справочник клиентов и БИН из текущей базы.",
        "Договоры, сроки действия и привязка к компании.",
        "Виды топлива, цены, скидки и способы поставки.",
        "Шаблоны документов и обязательные комплекты.",
        "Исторические продажи для реальной аналитики и прогноза.",
        "Примеры реальных обращений, возвратов, продлений и тендеров.",
    ]
    for item in data_items:
        add_list_item(doc, item, bullet_num_id)

    add_section_heading(doc, "11. Рекомендуемая дорожная карта")
    roadmap_rows = [
        ["Этап 1", "Зафиксировать бизнес-правила", "SLA, владельцы ролей, обязательные поля, документы, переходы", "1–2 недели"],
        ["Этап 2", "Сервер и база данных", "Backend API, PostgreSQL, серверная авторизация, миграции, резервные копии", "3–5 недель"],
        ["Этап 3", "Файлы и безопасность", "Хранилище документов, права, версии, журнал доступа, HTTPS", "2–3 недели"],
        ["Этап 4", "Омниканальные входящие", "WhatsApp, телефония, email, webhooks, автоматическое распределение", "3–6 недель"],
        ["Этап 5", "Реальная интеграция 1С", "Счета, оплаты, реализации, задолженность, ЭСФ и акты", "4–8 недель"],
        ["Этап 6", "Реальные сообщения", "Серверный чат, группы, доставка, архив, вложения", "2–4 недели"],
        ["Этап 7", "ИИ и расширенная аналитика", "AI API, безопасный контекст, прогноз продаж, рекомендации", "3–6 недель"],
        ["Этап 8", "Пилот и запуск", "Тестовые данные, обучение, исправления, мониторинг и регламенты", "2–4 недели"],
    ]
    add_table(doc, ["Этап", "Цель", "Результат", "Оценка"], roadmap_rows, [900, 2100, 5100, 1260], font_size=8.3)
    add_callout(
        doc,
        "Рекомендация",
        "Не начинать реальные интеграции до утверждения SLA, владельцев ролей и обязательных данных. "
        "Иначе автоматизация закрепит неясные процессы и потребует дорогой переделки.",
        "blue",
    )

    add_heading(doc, "Приоритет ближайшего цикла", 2)
    priorities = [
        "Утвердить SLA менеджеров обслуживания и контрольные показатели Дианы.",
        "Утвердить кассу, бухгалтера по банку и правила банковского возврата.",
        "Подготовить серверную архитектуру и модель базы данных.",
        "Выбрать первый реальный канал интеграции — рекомендуется WhatsApp или корпоративная почта.",
        "Согласовать контракт интеграции с 1С на примере одного заказа.",
    ]
    priority_num_id = duplicate_numbering_instance(doc, decimal_num_id)
    for item in priorities:
        add_list_item(doc, item, priority_num_id)

    add_section_heading(doc, "12. Контрольный список следующего этапа")
    checklist = [
        "Назначен владелец каждого процесса.",
        "Утверждены SLA и правила эскалации.",
        "Утверждены роли кассы и бухгалтера по банку.",
        "Утверждены обязательные поля и документы по стадиям.",
        "Подготовлена схема базы данных.",
        "Выбран backend-стек и способ развёртывания.",
        "Выбран провайдер WhatsApp и телефонии.",
        "Получено описание обмена с 1С.",
        "Определено файловое хранилище.",
        "Определён пилотный контур и группа пользователей.",
        "Подготовлены реальные тестовые кейсы для приёмки.",
    ]
    for item in checklist:
        add_list_item(doc, f"☐ {item}", bullet_num_id)

    add_heading(doc, "Предлагаемый следующий рабочий документ", 2)
    add_body(
        doc,
        "После утверждения этой карты рекомендуется подготовить отдельный документ «Регламент ролей, SLA и эскалаций GAZOIL». "
        "Он должен содержать точные сроки, ответственных, заместителей, основания для возврата карточки и правила уведомления руководства.",
    )

    add_section_heading(doc, "Приложение. Стартовые учётные записи")
    add_callout(
        doc,
        "Безопасность",
        "Стартовый пароль Aa123456 используется только для демонстрации. В рабочей системе временный пароль должен выдаваться индивидуально, "
        "с обязательной сменой при первом входе и серверным хэшированием.",
        "red",
    )
    accounts = [
        ["Мади", "Мади / комдир", "Aa123456"],
        ["Рашид", "Рашид / RR", "Aa123456"],
        ["Арсен", "Арсен / AB", "Aa123456"],
        ["Диана", "Диана", "Aa123456"],
        ["Ельжан", "Ельжан", "Aa123456"],
        ["Жанара", "Жанара", "Aa123456"],
        ["Екатерина", "Екатерина", "Aa123456"],
        ["Ольга", "Ольга", "Aa123456"],
        ["Павел", "Павел", "Aa123456"],
        ["Ильяс", "Ильяс", "Aa123456"],
        ["Татьяна", "Татьяна", "Aa123456"],
        ["Айнура", "Айнура", "Aa123456"],
        ["Учредитель", "Учредитель", "Aa123456"],
    ]
    add_table(doc, ["Пользователь", "Логин", "Стартовый пароль"], accounts, [3000, 3360, 3000], font_size=9.2)

    add_heading(doc, "Итог", 1)
    add_body(
        doc,
        "Проект уже имеет связную бизнес-логику и достаточно экранов для демонстрации будущей системы. "
        "Следующий качественный скачок — не добавление ещё одного экрана, а перевод прототипа на серверную архитектуру "
        "после утверждения SLA, владельцев ролей и точных правил процессов.",
    )
    add_callout(
        doc,
        "Рекомендуемое действие",
        "Начать с совместного утверждения регламента ролей и сроков, затем проектировать базу данных и только после этого подключать реальные каналы и 1С.",
        "green",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
